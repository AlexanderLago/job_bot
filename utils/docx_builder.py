# utils/docx_builder.py — Build an ATS-optimized .docx resume

import io
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── Helpers ────────────────────────────────────────────────────────────────────

def _add_horizontal_rule(doc):
    """Add a thin horizontal line below a paragraph."""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(2)
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "4F46E5")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return para


def _set_para_spacing(para, before=0, after=2):
    para.paragraph_format.space_before = Pt(before)
    para.paragraph_format.space_after = Pt(after)


def _section_header(doc, title: str):
    """Add an ATS-friendly section header with a rule underneath."""
    para = doc.add_paragraph()
    _set_para_spacing(para, before=8, after=0)
    run = para.add_run(title.upper())
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x4F, 0x46, 0xE5)
    _add_horizontal_rule(doc)


def _bullet(doc, text: str):
    """Add a bullet point paragraph."""
    para = doc.add_paragraph(style="List Bullet")
    _set_para_spacing(para, before=0, after=1)
    run = para.add_run(text)
    run.font.size = Pt(10)
    run.font.name = "Calibri"


def _normal(doc, text: str, bold=False, size=10, color=None, align=WD_ALIGN_PARAGRAPH.LEFT):
    para = doc.add_paragraph()
    _set_para_spacing(para, before=0, after=1)
    para.alignment = align
    run = para.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Calibri"
    if color:
        run.font.color.rgb = color
    return para


# ── Main builder ──────────────────────────────────────────────────────────────

def build_docx(data: dict) -> bytes:
    """
    Build an ATS-optimized .docx from structured resume data.
    Returns bytes ready for download.
    """
    doc = Document()

    # Page margins (0.75 inch all around — ATS-safe range)
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    # Set default font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    # ── Name ──────────────────────────────────────────────────────────────────
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_para_spacing(name_para, before=0, after=2)
    name_run = name_para.add_run(data.get("name", ""))
    name_run.bold = True
    name_run.font.size = Pt(18)
    name_run.font.name = "Calibri"

    # ── Contact line ──────────────────────────────────────────────────────────
    contact_parts = [
        data.get("phone", ""),
        data.get("email", ""),
        data.get("location", ""),
        data.get("linkedin", ""),
        data.get("website", ""),
    ]
    contact_line = "  |  ".join(p for p in contact_parts if p)
    if contact_line:
        contact_para = doc.add_paragraph()
        contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_para_spacing(contact_para, before=0, after=6)
        contact_run = contact_para.add_run(contact_line)
        contact_run.font.size = Pt(9)
        contact_run.font.name = "Calibri"
        contact_run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    # ── Summary ───────────────────────────────────────────────────────────────
    summary = data.get("summary", "")
    if summary:
        _section_header(doc, "Professional Summary")
        para = doc.add_paragraph()
        _set_para_spacing(para, before=2, after=4)
        run = para.add_run(summary)
        run.font.size = Pt(10)
        run.font.name = "Calibri"

    # ── Experience ────────────────────────────────────────────────────────────
    experience = data.get("experience", [])
    if experience:
        _section_header(doc, "Experience")
        for job in experience:
            # Title + Company line
            title_para = doc.add_paragraph()
            _set_para_spacing(title_para, before=4, after=0)
            title_run = title_para.add_run(job.get("title", ""))
            title_run.bold = True
            title_run.font.size = Pt(11)
            title_run.font.name = "Calibri"

            # Company + location + dates (same paragraph, right-aligned dates via tab)
            meta_para = doc.add_paragraph()
            _set_para_spacing(meta_para, before=0, after=1)
            company_str = job.get("company", "")
            location_str = job.get("location", "")
            dates_str = job.get("dates", "")
            left_text = f"{company_str}" + (f"  —  {location_str}" if location_str else "")
            meta_run_left = meta_para.add_run(left_text)
            meta_run_left.italic = True
            meta_run_left.font.size = Pt(10)
            meta_run_left.font.name = "Calibri"
            meta_run_left.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
            if dates_str:
                meta_para.add_run("   ")
                date_run = meta_para.add_run(dates_str)
                date_run.font.size = Pt(10)
                date_run.font.name = "Calibri"
                date_run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

            for bullet in job.get("bullets", []):
                _bullet(doc, bullet)

    # ── Education ─────────────────────────────────────────────────────────────
    education = data.get("education", [])
    if education:
        _section_header(doc, "Education")
        for edu in education:
            degree_para = doc.add_paragraph()
            _set_para_spacing(degree_para, before=4, after=0)
            d_run = degree_para.add_run(edu.get("degree", ""))
            d_run.bold = True
            d_run.font.size = Pt(11)
            d_run.font.name = "Calibri"

            school_para = doc.add_paragraph()
            _set_para_spacing(school_para, before=0, after=1)
            school_text = edu.get("school", "")
            loc_text = edu.get("location", "")
            dates_text = edu.get("dates", "")
            left = f"{school_text}" + (f"  —  {loc_text}" if loc_text else "")
            s_run = school_para.add_run(left)
            s_run.italic = True
            s_run.font.size = Pt(10)
            s_run.font.name = "Calibri"
            s_run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
            if dates_text:
                school_para.add_run("   ")
                school_para.add_run(dates_text).font.size = Pt(10)

            details = edu.get("details", "")
            if details:
                det_para = doc.add_paragraph()
                _set_para_spacing(det_para, before=0, after=1)
                det_para.add_run(details).font.size = Pt(10)

    # ── Skills ────────────────────────────────────────────────────────────────
    skills = data.get("skills", [])
    if skills:
        _section_header(doc, "Skills")
        # Group skills into rows of ~4 for readability
        chunk_size = 4
        rows = [skills[i:i+chunk_size] for i in range(0, len(skills), chunk_size)]
        for row in rows:
            para = doc.add_paragraph()
            _set_para_spacing(para, before=1, after=1)
            run = para.add_run("  •  ".join(row))
            run.font.size = Pt(10)
            run.font.name = "Calibri"

    # ── Certifications ────────────────────────────────────────────────────────
    certs = data.get("certifications", [])
    if certs:
        _section_header(doc, "Certifications")
        for cert in certs:
            _bullet(doc, cert)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
