# utils/log_builder.py — Export the application log as DOCX or CSV

from __future__ import annotations

import csv
import io
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


_COLUMNS = ["Date", "Job Title", "Company", "Location", "Work Type", "Fit %"]
_KEYS = ["date", "job_title", "company", "location", "work_type", "fit_pct"]


def build_log_docx(app_log: list[dict]) -> bytes:
    """
    Build a Word document containing a styled application-tracking table.
    Returns bytes ready for download.
    """
    doc = Document()

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Job Application Log")
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = "Calibri"
    run.font.color.rgb = RGBColor(0x4F, 0x46, 0xE5)

    doc.add_paragraph()  # spacer

    table = doc.add_table(rows=1, cols=len(_COLUMNS))
    table.style = "Table Grid"

    # Header row
    hdr_cells = table.rows[0].cells
    for i, col_name in enumerate(_COLUMNS):
        hdr_cells[i].text = col_name
        run = hdr_cells[i].paragraphs[0].runs[0]
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = "Calibri"
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        # Dark background for header cells
        tc_pr = hdr_cells[i]._tc.get_or_add_tcPr()
        shd = _make_shading("4F46E5")
        tc_pr.append(shd)

    # Data rows
    for entry in app_log:
        row_cells = table.add_row().cells
        for i, key in enumerate(_KEYS):
            val = entry.get(key, "")
            if val is None:
                val = ""
            row_cells[i].text = str(val)
            run = row_cells[i].paragraphs[0].runs[0] if row_cells[i].paragraphs[0].runs else row_cells[i].paragraphs[0].add_run(str(val))
            run.font.size = Pt(10)
            run.font.name = "Calibri"

    # Auto-fit columns (approximate)
    for col in table.columns:
        for cell in col.cells:
            cell.paragraphs[0].paragraph_format.space_before = Pt(2)
            cell.paragraphs[0].paragraph_format.space_after = Pt(2)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def build_log_csv(app_log: list[dict]) -> str:
    """
    Build a CSV string from the application log.
    Returns a UTF-8 string suitable for st.download_button.
    """
    output = io.StringIO()
    writer = csv.DictWriter(
        output,
        fieldnames=_KEYS,
        extrasaction="ignore",
    )
    writer.writeheader()
    for entry in app_log:
        writer.writerow({k: entry.get(k, "") or "" for k in _KEYS})
    return output.getvalue()


def _make_shading(hex_color: str):
    """Create an OxmlElement for cell background shading."""
    from docx.oxml import OxmlElement
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    return shd
