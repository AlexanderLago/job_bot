# utils/resume_parser.py — Extract plain text from uploaded resume files

import io


def parse_resume(file_bytes: bytes, filename: str) -> str:
    """
    Extract plain text from a resume file.
    Supports: .docx, .pdf, .txt
    Returns the full text content as a string.
    """
    ext = filename.lower().rsplit(".", 1)[-1]

    if ext == "docx":
        return _parse_docx(file_bytes)
    elif ext == "pdf":
        return _parse_pdf(file_bytes)
    elif ext == "txt":
        return file_bytes.decode("utf-8", errors="ignore")
    else:
        raise ValueError(f"Unsupported file type: .{ext}")


def _parse_docx(file_bytes: bytes) -> str:
    from docx import Document
    doc = Document(io.BytesIO(file_bytes))
    lines = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            lines.append(text)
    # Also grab text from tables (some resumes use them)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text and text not in lines:
                    lines.append(text)
    return "\n".join(lines)


def _parse_pdf(file_bytes: bytes) -> str:
    try:
        import pdfplumber
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            pages = []
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    pages.append(text)
        return "\n".join(pages)
    except Exception as e:
        raise RuntimeError(f"Could not read PDF: {e}. Try converting to .docx or .txt.")
